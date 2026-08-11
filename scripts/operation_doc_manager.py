#!/usr/bin/env python3
"""
operation_doc_manager.py - 操作文档管理工具（独立于 versions）

支持功能：
- init: 初始化某功能的操作文档目录（含 screenshots 子目录）
- create: 从模板创建 Markdown 文档骨架
- generate-docx: 将 Markdown 转为 docx 格式（严格对齐 操作文档编写参考模版.docx 模板）
- review: 生成操作文档评审报告
- validate: 验证操作文档完整性
- list: 列出所有操作文档

输出目录结构：
    operation_docs/
        templates/        # 模板（md + docx）
        markdown/         # Markdown 源文件
        docx/             # docx 渲染版本
        screenshots/
            {功能名}/     # 各功能的截图

用法：
    python3 scripts/operation_doc_manager.py init --feature 议价策略
    python3 scripts/operation_doc_manager.py create --feature 议价策略 [--author ChaosXiong]
    python3 scripts/operation_doc_manager.py generate-docx --feature 议价策略
    python3 scripts/operation_doc_manager.py review --feature 议价策略
    python3 scripts/operation_doc_manager.py validate --feature 议价策略
    python3 scripts/operation_doc_manager.py list
"""

import argparse
import re
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("[错误] 缺少 python-docx，请运行: pip3 install python-docx")
    sys.exit(1)


PROJECT_ROOT = Path(__file__).parent.parent.resolve()
OP_DOCS_ROOT = PROJECT_ROOT / "operation_docs"
TEMPLATES_DIR = OP_DOCS_ROOT / "templates"
MARKDOWN_DIR = OP_DOCS_ROOT / "markdown"
DOCX_DIR = OP_DOCS_ROOT / "docx"
SCREENSHOTS_DIR = OP_DOCS_ROOT / "screenshots"

MD_TEMPLATE = TEMPLATES_DIR / "operation_doc_template.md"


def ensure_dirs():
    """确保所有必要目录存在。"""
    for d in [OP_DOCS_ROOT, TEMPLATES_DIR, MARKDOWN_DIR, DOCX_DIR, SCREENSHOTS_DIR]:
        d.mkdir(parents=True, exist_ok=True)


def cmd_init(args):
    """初始化某功能的操作文档目录。"""
    ensure_dirs()
    feature = args.feature
    feature_ss_dir = SCREENSHOTS_DIR / feature
    feature_ss_dir.mkdir(exist_ok=True)

    print(f"[OK] 操作文档目录已初始化")
    print(f"  - 文档根目录: {OP_DOCS_ROOT}")
    print(f"  - 功能截图目录: {feature_ss_dir}")
    print(f"  - Markdown 目录: {MARKDOWN_DIR}")
    print(f"  - docx 目录: {DOCX_DIR}")


def cmd_create(args):
    """从模板创建 Markdown 文档骨架。"""
    ensure_dirs()
    feature = args.feature
    today = datetime.now().strftime("%Y-%m-%d")
    author = args.author or "VitoMa"

    if not MD_TEMPLATE.exists():
        print(f"[错误] 模板不存在: {MD_TEMPLATE}")
        sys.exit(1)

    content = MD_TEMPLATE.read_text(encoding="utf-8")
    content = content.replace("{功能名}", feature)
    content = content.replace("{版本号}", args.version or "v1.0")
    content = content.replace("{日期}", today)
    content = content.replace("{作者}", author)
    content = content.replace("{状态}", "draft")

    doc_path = MARKDOWN_DIR / f"{feature}.md"
    doc_path.write_text(content, encoding="utf-8")

    # 同时初始化截图目录
    (SCREENSHOTS_DIR / feature).mkdir(exist_ok=True)

    print(f"[OK] Markdown 骨架已创建: {doc_path}")
    print(f"  - 功能名: {feature}")
    print(f"  - 作者: {author}")
    print(f"  - 截图目录: {SCREENSHOTS_DIR / feature}")


# ============================================================
# Markdown 解析与 docx 生成
# ============================================================

class MarkdownParser:
    """简易 Markdown 解析器：解析 Heading、段落、表格、图片、列表。"""

    def __init__(self, text: str):
        self.lines = text.split("\n")
        self.pos = 0
        self.blocks = []

    def parse(self):
        while self.pos < len(self.lines):
            line = self.lines[self.pos]

            # 跳过 frontmatter
            if self.pos == 0 and line.strip() == "---":
                self._skip_frontmatter()
                continue

            stripped = line.strip()
            if not stripped:
                self.pos += 1
                continue

            # 水平分隔线
            if stripped == "---":
                self.blocks.append({"type": "hr"})
                self.pos += 1
                continue

            # 标题
            m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if m:
                level = len(m.group(1))
                text = m.group(2)
                self.blocks.append({"type": "heading", "level": level, "text": text})
                self.pos += 1
                continue

            # 引用块
            if stripped.startswith(">"):
                quote = stripped[1:].strip()
                self.blocks.append({"type": "quote", "text": quote})
                self.pos += 1
                continue

            # 表格
            if "|" in line and self.pos + 1 < len(self.lines) and re.match(r"^\s*\|?[\s\-\|:]+\|?\s*$", self.lines[self.pos + 1]):
                table = self._parse_table()
                self.blocks.append(table)
                continue

            # 图片
            img_match = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", stripped)
            if img_match:
                self.blocks.append({"type": "image", "alt": img_match.group(1), "src": img_match.group(2)})
                self.pos += 1
                continue

            # 列表
            if re.match(r"^[\-\*]\s+", stripped):
                items = self._parse_list()
                self.blocks.append({"type": "list", "items": items})
                continue

            # 普通段落
            self.blocks.append({"type": "paragraph", "text": stripped})
            self.pos += 1

        return self.blocks

    def _skip_frontmatter(self):
        self.pos += 1
        while self.pos < len(self.lines) and self.lines[self.pos].strip() != "---":
            self.pos += 1
        self.pos += 1

    def _parse_table(self):
        header_line = self.lines[self.pos]
        headers = [c.strip() for c in header_line.strip().strip("|").split("|")]
        self.pos += 2  # 跳过分隔线

        rows = []
        while self.pos < len(self.lines):
            line = self.lines[self.pos]
            if "|" not in line or not line.strip():
                break
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(cells)
            self.pos += 1

        return {"type": "table", "headers": headers, "rows": rows}

    def _parse_list(self):
        items = []
        while self.pos < len(self.lines):
            line = self.lines[self.pos]
            m = re.match(r"^[\-\*]\s+(.+)$", line.strip())
            if not m:
                break
            items.append(m.group(1))
            self.pos += 1
        return items


def _set_cell_text(cell, text, bold=False, font_size=10):
    """设置表格单元格的文字。"""
    cell.text = ""
    para = cell.paragraphs[0]
    # 支持 <br> 换行
    parts = text.split("<br>")
    for i, part in enumerate(parts):
        if i > 0:
            para.add_run().add_break()
        run = para.add_run(part)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def _add_paragraph(doc, text, bold=False, size=11, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def _add_heading_styled(doc, text, level):
    """添加标题，模仿 操作文档编写参考模版.docx 的样式（粗体、字号梯度）。"""
    sizes = {1: 18, 2: 16, 3: 14, 4: 12, 5: 11, 6: 11}
    size = sizes.get(level, 11)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def md_to_docx(md_path: Path, docx_path: Path, screenshots_root: Path):
    """将 Markdown 文件转为 docx。"""
    text = md_path.read_text(encoding="utf-8")
    blocks = MarkdownParser(text).parse()

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            _add_heading_styled(doc, block["text"], block["level"])

        elif btype == "paragraph":
            _add_paragraph(doc, block["text"])

        elif btype == "quote":
            _add_paragraph(doc, block["text"], size=10)

        elif btype == "list":
            for item in block["items"]:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(item)
                run.font.size = Pt(11)
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        elif btype == "image":
            src = block["src"]
            # 解析路径：处理 ../screenshots/xxx
            if src.startswith("../"):
                img_path = (md_path.parent / src).resolve()
            else:
                img_path = (md_path.parent / src).resolve()

            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Cm(14))
                except Exception as e:
                    _add_paragraph(doc, f"[图片加载失败: {src} - {e}]", size=9)
            else:
                _add_paragraph(doc, f"[截图占位: {src}]", size=9)

        elif btype == "table":
            headers = block["headers"]
            rows = block["rows"]
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"

            # 表头
            for ci, h in enumerate(headers):
                _set_cell_text(table.rows[0].cells[ci], h, bold=True, font_size=10)

            # 数据行
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    if ci >= len(headers):
                        continue
                    # 处理图片引用
                    img_match = re.match(r"!\[.*?\]\((.*?)\)", cell_text)
                    if img_match:
                        img_src = img_match.group(1)
                        if img_src.startswith("../"):
                            img_path = (md_path.parent / img_src).resolve()
                        else:
                            img_path = (md_path.parent / img_src).resolve()
                        if img_path.exists():
                            cell = table.rows[ri + 1].cells[ci]
                            cell.text = ""
                            try:
                                cell.paragraphs[0].add_run().add_picture(str(img_path), width=Cm(4))
                            except Exception:
                                _set_cell_text(cell, "[图片占位]", font_size=9)
                        else:
                            _set_cell_text(table.rows[ri + 1].cells[ci], "[截图占位]", font_size=9)
                    else:
                        _set_cell_text(table.rows[ri + 1].cells[ci], cell_text, font_size=10)

        elif btype == "hr":
            doc.add_paragraph("─" * 40)

    doc.save(str(docx_path))


def cmd_generate_docx(args):
    """将 Markdown 转为 docx。"""
    ensure_dirs()
    feature = args.feature
    md_path = MARKDOWN_DIR / f"{feature}.md"
    if not md_path.exists():
        print(f"[错误] Markdown 文档不存在: {md_path}")
        print(f"  请先运行: python3 scripts/operation_doc_manager.py create --feature {feature}")
        sys.exit(1)

    docx_path = DOCX_DIR / f"{feature}.docx"
    md_to_docx(md_path, docx_path, SCREENSHOTS_DIR / feature)

    print(f"[OK] docx 已生成: {docx_path}")
    print(f"  - 源 Markdown: {md_path}")
    print(f"  - 输出 docx: {docx_path}")


def cmd_review(args):
    """生成操作文档评审报告。"""
    ensure_dirs()
    feature = args.feature
    md_path = MARKDOWN_DIR / f"{feature}.md"
    if not md_path.exists():
        print(f"[错误] 文档不存在: {md_path}")
        sys.exit(1)

    content = md_path.read_text(encoding="utf-8")
    scores = {}
    issues = []

    # 1. 结构完整性（30分）
    structure_score = 30
    required_sections = [
        "## 一、功能介绍",
        "## 二、",
        "## 四、常见问题",
    ]
    for section in required_sections:
        if section not in content:
            structure_score -= 10
            issues.append(f"缺少必要章节: {section}")
    scores["结构完整性"] = max(0, structure_score)

    # 2. 截图覆盖率（30分）
    img_count = len(re.findall(r"!\[.*?\]\(.*?\)", content))
    step_headers = len(re.findall(r"^### \d+、", content, re.MULTILINE))
    coverage = (img_count / step_headers) if step_headers else 0
    screenshot_score = min(30, int(coverage * 30)) if step_headers else 0
    if step_headers and coverage < 0.8:
        issues.append(f"截图覆盖率偏低: {img_count}/{step_headers} 步骤有截图 ({coverage*100:.0f}%)")
    scores["截图覆盖率"] = screenshot_score

    # 3. 内容准确性（20分）
    accuracy_score = 20
    placeholder_count = len(re.findall(r"\{[^{}]+\}", content))
    if placeholder_count > 5:
        accuracy_score -= min(15, placeholder_count)
        issues.append(f"文档中仍有 {placeholder_count} 处未填充的占位符")
    scores["内容准确性"] = max(0, accuracy_score)

    # 4. docx 同步（20分）
    docx_score = 0
    docx_path = DOCX_DIR / f"{feature}.docx"
    if docx_path.exists():
        docx_score = 20
    else:
        issues.append(f"docx 文件未生成: {docx_path}")
    scores["docx同步"] = docx_score

    total = sum(scores.values())

    report = f"""---
title: 操作文档评审报告
doc: {feature}
date: {datetime.now().strftime("%Y-%m-%d %H:%M")}
---

# 操作文档评审报告

| 项目 | 内容 |
|------|------|
| 评审文档 | {feature} |
| Markdown | {md_path} |
| docx | {docx_path} |
| 综合评分 | {total}/100 |
| 评审结果 | {'通过' if total >= 85 else '不通过'} |

## 评分详情

| 维度 | 满分 | 得分 |
|------|------|------|
| 结构完整性 | 30 | {scores['结构完整性']} |
| 截图覆盖率 | 30 | {scores['截图覆盖率']} |
| 内容准确性 | 20 | {scores['内容准确性']} |
| docx 同步 | 20 | {scores['docx同步']} |

## 统计

- 操作章节数: {step_headers}
- 截图引用数: {img_count}
- 截图覆盖率: {coverage*100:.0f}%
- 占位符残留: {placeholder_count}

## 问题清单

"""
    if issues:
        for i, issue in enumerate(issues, 1):
            report += f"{i}. {issue}\n"
    else:
        report += "无问题，文档质量良好。\n"

    report += f"\n## 结论\n\n{'评审通过，文档可发布。' if total >= 85 else '评审不通过，请修复上述问题后重新评审。'}\n"

    report_path = OP_DOCS_ROOT / f"{feature}-评审报告.md"
    report_path.write_text(report, encoding="utf-8")

    print(f"[OK] 评审报告已生成: {report_path}")
    print(f"  - 综合评分: {total}/100")
    print(f"  - 评审结果: {'通过' if total >= 85 else '不通过'}")
    if issues:
        print(f"  - 发现问题: {len(issues)} 项")


def cmd_validate(args):
    """验证操作文档完整性。"""
    ensure_dirs()
    feature = args.feature
    md_path = MARKDOWN_DIR / f"{feature}.md"
    docx_path = DOCX_DIR / f"{feature}.docx"
    ss_dir = SCREENSHOTS_DIR / feature

    if not md_path.exists():
        print(f"[错误] 文档不存在: {md_path}")
        sys.exit(1)

    content = md_path.read_text(encoding="utf-8")

    checks = {
        "Markdown 文件存在": md_path.exists(),
        "docx 文件存在": docx_path.exists(),
        "截图目录存在": ss_dir.exists(),
        "包含 一、功能介绍": "## 一、功能介绍" in content,
        "包含 二、操作说明": "## 二、" in content,
        "包含 四、常见问题": "## 四、常见问题" in content,
    }

    img_refs = re.findall(r"!\[.*?\]\((.*?)\)", content)
    missing_imgs = []
    for ref in img_refs:
        if ref.startswith("../"):
            img_path = (md_path.parent / ref).resolve()
        else:
            img_path = (md_path.parent / ref).resolve()
        if not img_path.exists():
            missing_imgs.append(ref)

    checks["所有截图文件存在"] = len(missing_imgs) == 0

    print(f"验证结果: {feature}")
    print("-" * 40)
    all_pass = True
    for check, passed in checks.items():
        status = "[OK]" if passed else "[FAIL]"
        print(f"  {status} {check}")
        if not passed:
            all_pass = False

    if missing_imgs:
        print(f"\n  缺失的截图（{len(missing_imgs)} 张）:")
        for img in missing_imgs[:10]:
            print(f"    - {img}")

    if all_pass:
        print(f"\n[OK] 所有检查项通过")
    else:
        print(f"\n[FAIL] 存在未通过项，请修复后重试")
        sys.exit(1)


def cmd_list(args):
    """列出所有操作文档。"""
    ensure_dirs()
    md_files = sorted(MARKDOWN_DIR.glob("*.md"))
    docx_files = sorted(DOCX_DIR.glob("*.docx"))

    print(f"操作文档清单（{OP_DOCS_ROOT}）")
    print("-" * 50)

    if md_files:
        print(f"\nMarkdown 文档 ({len(md_files)}):")
        for f in md_files:
            print(f"  - {f.name}")

    if docx_files:
        print(f"\ndocx 文档 ({len(docx_files)}):")
        for f in docx_files:
            print(f"  - {f.name}")

    if SCREENSHOTS_DIR.exists():
        ss_subdirs = [d for d in SCREENSHOTS_DIR.iterdir() if d.is_dir()]
        if ss_subdirs:
            print(f"\n截图目录 ({len(ss_subdirs)}):")
            for d in sorted(ss_subdirs):
                pngs = list(d.glob("*.png"))
                print(f"  - {d.name}/ ({len(pngs)} 张)")

    if not md_files and not docx_files:
        print("  暂无文档")


def main():
    parser = argparse.ArgumentParser(
        description="操作文档管理工具（独立于 versions）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python3 scripts/operation_doc_manager.py init --feature 议价策略
  python3 scripts/operation_doc_manager.py create --feature 议价策略 --author ChaosXiong
  python3 scripts/operation_doc_manager.py generate-docx --feature 议价策略
  python3 scripts/operation_doc_manager.py review --feature 议价策略
  python3 scripts/operation_doc_manager.py validate --feature 议价策略
  python3 scripts/operation_doc_manager.py list
        """,
    )

    sub = parser.add_subparsers(dest="command", help="子命令")

    p_init = sub.add_parser("init", help="初始化某功能的操作文档目录")
    p_init.add_argument("--feature", required=True, help="功能名称")

    p_create = sub.add_parser("create", help="从模板创建 Markdown 骨架")
    p_create.add_argument("--feature", required=True, help="功能名称")
    p_create.add_argument("--author", help="作者")
    p_create.add_argument("--version", help="文档版本号（默认 v1.0）")

    p_docx = sub.add_parser("generate-docx", help="Markdown 转 docx")
    p_docx.add_argument("--feature", required=True, help="功能名称")

    p_review = sub.add_parser("review", help="生成评审报告")
    p_review.add_argument("--feature", required=True, help="功能名称")

    p_validate = sub.add_parser("validate", help="验证文档完整性")
    p_validate.add_argument("--feature", required=True, help="功能名称")

    sub.add_parser("list", help="列出所有操作文档")

    args = parser.parse_args()

    if args.command is None:
        parser.print_help()
        sys.exit(1)

    commands = {
        "init": cmd_init,
        "create": cmd_create,
        "generate-docx": cmd_generate_docx,
        "review": cmd_review,
        "validate": cmd_validate,
        "list": cmd_list,
    }

    commands[args.command](args)


if __name__ == "__main__":
    main()
